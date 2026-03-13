# -*- coding: utf-8 -*-
"""
Generate the kid-friendly assembly manual as a Word document.
Run: python3 generate_word.py
Output: 小朋友的组装手册.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Global styles ──

style = doc.styles['Normal']
style.font.name = '\u5fae\u8f6f\u96c5\u9ed1'  # 微软雅黑
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

FONT_CN = '\u5fae\u8f6f\u96c5\u9ed1'
BLUE = RGBColor(0x1A, 0x56, 0xDB)
GREEN = RGBColor(0x00, 0x80, 0x00)
GOLD = RGBColor(0xE8, 0x9D, 0x0E)
RED_WARN = RGBColor(0xE8, 0x4D, 0x0E)
GREY = RGBColor(0x33, 0x33, 0x33)

for lvl in range(1, 4):
    hs = doc.styles['Heading %d' % lvl]
    hs.font.name = FONT_CN
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    hs.font.bold = True
    hs.font.color.rgb = BLUE if lvl <= 2 else GREY
    hs.font.size = {1: Pt(22), 2: Pt(16), 3: Pt(13)}[lvl]
    hs.paragraph_format.space_before = Pt(18 if lvl == 1 else 14)
    hs.paragraph_format.space_after = Pt(6)

for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

# ── Helpers ──

def _run(p, text, bold=False, italic=False, size=None, color=None, font_name=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    fn = font_name or FONT_CN
    r.font.name = fn
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    return r

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text, bold=False, italic=False, size=None, color=None, align=None):
    para = doc.add_paragraph()
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(para, text, bold=bold, italic=italic, size=size, color=color)
    return para

def code(text):
    for line in text.strip().split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.1
        r = para.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(9.5)
        r.font.color.rgb = GREY

def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    _run(para, text, size=10.5)

def checklist(items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        _run(para, '\u2610  ' + item, size=10.5)

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for rp in c.paragraphs:
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in rp.runs:
                rn.bold = True
                rn.font.name = FONT_CN
                rn.font.size = Pt(10)
                rn.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = val
            for rp in c.paragraphs:
                for rn in rp.runs:
                    rn.font.name = FONT_CN
                    rn.font.size = Pt(10)
                    rn.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

def warn(text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.text = ''
    for line in text.strip().split('\n'):
        cp = cell.add_paragraph(line)
        cp.paragraph_format.space_after = Pt(2)
        for rn in cp.runs:
            rn.font.name = FONT_CN
            rn.font.size = Pt(10.5)
            rn.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:%s' % edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'E84D0E')
        borders.append(e)
    tc_pr.append(borders)
    cell.paragraphs[0].clear()

def star(text):
    p(text, bold=True, size=13, color=GOLD, align='center')

def sep():
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run('\u2500' * 50)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(8)

LQ = '\u300c'  # 「
RQ = '\u300d'  # 」

# ══════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()
p('\U0001F697', size=48, align='center')
p('\u6211\u7684\u667a\u80fd\u6e05\u6d01\u5c0f\u8f66', bold=True, size=28, color=BLUE, align='center')  # 我的智能清洁小车
p('\u7ec4\u88c5\u5927\u5192\u9669\uff01', bold=True, size=22, color=BLUE, align='center')  # 组装大冒险！

doc.add_paragraph()
p('\u4e00\u6b65\u4e00\u6b65\uff0c\u628a\u4e00\u5806\u96f6\u4ef6\u53d8\u6210\u4e00\u8f86', size=14, align='center')
p('\u4f1a\u8bf4\u8bdd\u3001\u4f1a\u626b\u5730\u7684\u667a\u80fd\u5c0f\u8f66\uff01', size=14, align='center')

doc.add_paragraph()
p('\u4e00\u5171 10 \u5173\uff0c\u6bcf\u5173\u5b8c\u6210\u8d34\u4e00\u9897 \u2b50', size=12, align='center')
p('\u5168\u90e8\u901a\u5173 = \u5c0f\u8f66\u8bde\u751f\uff01', bold=True, size=12, align='center')

for _ in range(4):
    doc.add_paragraph()
p('\u5c0f\u8f66\u5de5\u7a0b\u5e08\uff1a_______________', size=12, align='center')
p('\u5f00\u5de5\u65e5\u671f\uff1a_______________', size=12, align='center')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  ADVENTURE MAP
# ══════════════════════════════════════════════════

h('\U0001F5FA\uFE0F \u5192\u9669\u5730\u56fe', level=1)
p('\u4ece\u8d77\u70b9\u51fa\u53d1\uff0c\u4e00\u5173\u4e00\u5173\u5f80\u4e0a\u722c\uff01\u6bcf\u901a\u8fc7\u4e00\u5173\uff0c\u5728\u53f3\u8fb9\u7684\u65b9\u6846\u91cc\u6253\u4e2a\u52fe \u2713')
doc.add_paragraph()

map_data = [
    ['\U0001F3C1', '\u7ec8\u70b9', '\u5c0f\u8f66\u5b8c\u6210\uff01\u4e0a\u53f0\u5c55\u793a\uff01', '\u2610'],
    ['10', '\u5f69\u6392\u6f14\u7ec3', '\u7ec3\u4e60\u5c55\u793a', '\u2610'],
    ['9', '\u7a7f\u8863\u670d', '\u88c5\u5916\u58f3\u3001\u53d8\u6f02\u4eae', '\u2610'],
    ['8', '\u88c5\u7535\u6c60', '\u62d4\u6389\u7ebf\uff0c\u81ea\u7531\u5954\u8dd1', '\u2610'],
    ['7', '\u88c5\u626b\u628a', '\u4f1a\u626b\u5730\u5566', '\u2610'],
    ['6', '\u88c5\u8033\u6735', '\u542c\u61c2\u4f60\u8bf4\u8bdd', '\u2610'],
    ['5', '\u88c5\u5634\u5df4', '\u4f1a\u8bf4\u8bdd\u5566\uff08\u5f55\u53f0\u8bcd\uff01\uff09', '\u2610'],
    ['4', '\u88c5\u96f7\u8fbe', '\u4e0d\u649e\u4e1c\u897f', '\u2610'],
    ['3', '\u88c5\u773c\u775b', '\u4e0d\u6389\u684c\u5b50\uff08\u6700\u91cd\u8981\uff01\uff09', '\u2610'],
    ['2', '\u88c5\u817f', '\u4f1a\u8d70\u8def\u5566', '\u2610'],
    ['1', '\u70b9\u4eae\u5c0f\u706f', '\u7b2c\u4e00\u6b21\u8ba9\u786c\u4ef6\u52a8\u8d77\u6765\uff01', '\u2610'],
    ['0', '\u5f00\u7bb1\u8ba4\u96f6\u4ef6', '\u8ba4\u8bc6\u6240\u6709\u96f6\u4ef6', '\u2610'],
    ['\U0001F680', '\u8d77\u70b9', '\u4f60\u5728\u8fd9\u91cc\uff01', '\u2610'],
]
table(['\u5173\u5361', '\u540d\u79f0', '\u76ee\u6807', '\u901a\u8fc7 \u2713'], map_data)

doc.add_page_break()

# ══════════════════════════════════════════════════
#  LEVEL 0
# ══════════════════════════════════════════════════

h('\U0001F530 \u7b2c 0 \u5173\uff1a\u5f00\u7bb1\u8ba4\u96f6\u4ef6', level=1)
p('\U0001F3AF \u4efb\u52a1\uff1a\u8ba4\u8bc6\u6240\u6709\u96f6\u4ef6\uff0c\u7ed9\u5b83\u4eec\u627e\u5230\u81ea\u5df1\u7684\u5bb6', bold=True, size=12)

h('\U0001F4E6 \u4f60\u9700\u8981', level=2)
bullet('\u521a\u4e70\u56de\u6765\u7684\u6240\u6709\u96f6\u4ef6\u5feb\u9012')
bullet('\u4e00\u5f20\u5927\u684c\u5b50')
bullet('\u4fbf\u5229\u8d34 + \u7b14')
bullet('\u624b\u673a\uff08\u62cd\u7167\u7528\uff09')

h('\u7b2c \u2460 \u6b65\uff1a\u62c6\u5feb\u9012\uff01', level=2)
p('\u628a\u6240\u6709\u5feb\u9012\u62c6\u5f00\uff0c\u96f6\u4ef6\u90fd\u5012\u51fa\u6765\u3002\u522b\u6025\uff0c\u5148\u522b\u62c6\u96f6\u4ef6\u7684\u5c0f\u888b\u5b50\uff01')

h('\u7b2c \u2461 \u6b65\uff1a\u7ed9\u96f6\u4ef6\u5206\u5bb6', level=2)
p('\u5728\u684c\u4e0a\u5212\u51fa 6 \u4e2a\u533a\u57df\uff0c\u7528\u4fbf\u5229\u8d34\u6807\u4e0a\u540d\u5b57\uff0c\u628a\u96f6\u4ef6\u653e\u5230\u5bf9\u5e94\u7684\u5bb6\uff1a')
table(
    ['\u5206\u7ec4', '\u96f6\u4ef6'],
    [
        ['\U0001F9E0 \u5927\u8111\u7684\u5bb6', 'ESP32\u3001TB6612 \u9a71\u52a8\u677f'],
        ['\U0001F9BF \u817f\u7684\u5bb6', 'N20 \u7535\u673a \xd72\u3001\u8f6e\u5b50 \xd72\u3001\u4e07\u5411\u7403\u8f6e'],
        ['\U0001F441 \u773c\u775b\u7684\u5bb6', 'TCRT5000 \xd73\u3001HC-SR04 \u8d85\u58f0\u6ce2'],
        ['\U0001F50A \u5634\u5df4+\u8033\u6735\u7684\u5bb6', 'DFPlayer\u3001\u5c0f\u55ad\u53ed\u3001SD\u5361\u3001LU-ASR01\u30011k\u03a9\u7535\u963b'],
        ['\U0001F9F9 \u626b\u628a\u7684\u5bb6', '130 \u7535\u673a\u3001\u6bdb\u5237\u30013010 \u98ce\u6247\u3001MOS\u7ba1 \xd72'],
        ['\U0001F50B \u7535\u6c60\u7684\u5bb6', '18650 \u7535\u6c60 \xd72\u3001\u7535\u6c60\u76d2\u3001\u964d\u538b\u6a21\u5757'],
        ['\U0001F4CE \u5c0f\u96f6\u4ef6', 'LED \xd72\u3001\u8702\u9e23\u5668\u3001\u6309\u94ae\u3001\u675c\u90a6\u7ebf\u3001\u87ba\u4e1d\u94dc\u67f1'],
    ]
)

h('\u7b2c \u2462 \u6b65\uff1a\u5bf9\u7167\u6e05\u5355\u6253\u52fe', level=2)
p('\u627e\u5230\u4e00\u4e2a\uff0c\u5c31\u6253\u4e2a \u2713\uff1a')

p('\u3010\u5927\u8111\u7ec4\u3011', bold=True)
checklist(['ESP32 \u5f00\u53d1\u677f \xd71\uff08\u957f\u5f97\u50cf\u5c0f\u540d\u7247\uff0c\u4e24\u8fb9\u6709\u5f88\u591a\u5c0f\u811a\uff09',
           'TB6612 \u7535\u673a\u9a71\u52a8\u677f \xd71\uff08\u6307\u7532\u76d6\u5927\u5c0f\u7684\u5c0f\u677f\u5b50\uff09'])

p('\u3010\u817f\u7ec4\u3011', bold=True)
checklist(['N20 \u51cf\u901f\u7535\u673a \xd72\uff08\u62c7\u6307\u5927\u5c0f\u7684\u5c0f\u9a6c\u8fbe\uff09',
           '\u6a61\u80f6\u8f6e\u5b50 \xd72',
           '\u4e07\u5411\u7403\u8f6e \xd71\uff08\u5f39\u73e0\u5927\u5c0f\u7684\u91d1\u5c5e\u7403\uff09'])

p('\u3010\u773c\u775b\u7ec4\u3011', bold=True)
checklist(['TCRT5000 \u4f20\u611f\u5668 \xd73\uff08\u770b\u684c\u5b50\u8fb9\u7f18\u7684\u5c0f\u773c\u775b\uff09',
           'HC-SR04 \u8d85\u58f0\u6ce2 \xd71\uff08\u4e24\u4e2a\u5927\u5706\u773c\u775b\u7684\u6a21\u5757\uff09'])

p('\u3010\u5634\u5df4+\u8033\u6735\u7ec4\u3011', bold=True)
checklist(['DFPlayer Mini \xd71\uff08\u8ff7\u4f60 MP3 \u64ad\u653e\u5668\uff09',
           '\u5c0f\u55ad\u53ed \xd71\uff08\u5706\u5706\u7684\u8584\u7247\uff09',
           'Micro SD \u5361 \xd71',
           'LU-ASR01 \xd71\uff08\u8bed\u97f3\u8bc6\u522b\u6a21\u5757\uff09',
           '1k\u03a9 \u7535\u963b \xd71\uff08\u5c0f\u5c0f\u7684\u6761\u72b6\u96f6\u4ef6\uff0c\u6709\u5f69\u8272\u73af\uff09'])

p('\u3010\u626b\u628a\u7ec4\u3011', bold=True)
checklist(['130 \u5c0f\u7535\u673a \xd71', '3010 \u5c0f\u98ce\u6247 \xd71', 'MOS \u7ba1\u6a21\u5757 \xd72',
           '\u6bdb\u5237\uff08\u6216 DIY \u6750\u6599\uff09'])

p('\u3010\u7535\u6c60\u7ec4\u3011', bold=True)
checklist(['18650 \u9502\u7535\u6c60 \xd72', '\u53cc\u8282\u7535\u6c60\u76d2 \xd71\uff08\u5e26\u5f00\u5173\uff09',
           'LM2596 \u964d\u538b\u6a21\u5757 \xd71'])

p('\u3010\u5c0f\u96f6\u4ef6\u7ec4\u3011', bold=True)
checklist(['\u7eff\u8272 LED \xd71', '\u7ea2\u8272 LED \xd71', '\u8702\u9e23\u5668 \xd71',
           '\u6309\u94ae\u5f00\u5173 \xd71', '\u675c\u90a6\u7ebf\u4e00\u628a',
           '\u6d1e\u6d1e\u677f \xd71', '\u87ba\u4e1d\u94dc\u67f1\u82e5\u5e72'])

h('\u7b2c \u2463 \u6b65\uff1a\u62cd\u5168\u5bb6\u798f\uff01\U0001F4F8', level=2)
p('\u628a\u6240\u6709\u96f6\u4ef6\u6446\u597d\uff0c\u62cd\u4e00\u5f20\u7167\u7247\u3002\u4ee5\u540e\u505a\u6d77\u62a5\u4f1a\u7528\u5230\uff01')

h('\u7b2c \u2464 \u6b65\uff1a\u8003\u8003\u81ea\u5df1', level=2)
p('\u6307\u7740\u6bcf\u4e2a\u96f6\u4ef6\uff0c\u8bf4\u51fa\u5b83\u5728\u5c0f\u8f66\u8eab\u4e0a\u662f\u4ec0\u4e48\uff1a')
table(
    ['\u96f6\u4ef6', '\u5b83\u662f\u5c0f\u8f66\u7684\u2026\u2026'],
    [
        ['ESP32', '\U0001F9E0 \u5927\u8111\uff08\u6307\u6325\u6240\u6709\u96f6\u4ef6\uff09'],
        ['N20 \u7535\u673a', '\U0001F9BF \u817f\uff08\u8ba9\u5c0f\u8f66\u8d70\u8def\uff09'],
        ['TCRT5000', '\U0001F441 \u773c\u775b\uff08\u671d\u4e0b\u770b\uff0c\u9632\u6389\u684c\u5b50\uff09'],
        ['HC-SR04', '\U0001F4E1 \u96f7\u8fbe\uff08\u671d\u524d\u770b\uff0c\u9632\u649e\u4e1c\u897f\uff09'],
        ['DFPlayer', '\U0001F50A \u5634\u5df4\uff08\u8bf4\u8bdd\u7528\uff09'],
        ['LU-ASR01', '\U0001F3A4 \u8033\u6735\uff08\u542c\u4f60\u8bf4\u8bdd\uff09'],
        ['130 \u7535\u673a', '\U0001F9F9 \u626b\u628a\uff08\u5e26\u6bdb\u5237\u65cb\u8f6c\uff09'],
        ['3010 \u98ce\u6247', '\U0001F32C\uFE0F \u5438\u5c18\u5668\uff08\u5438\u788e\u5c51\uff09'],
        ['18650 \u7535\u6c60', '\U0001F50B \u98df\u7269\uff08\u7ed9\u5927\u5bb6\u4f9b\u7535\uff09'],
    ]
)

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u6240\u6709\u96f6\u4ef6\u90fd\u627e\u5230\u4e86\uff0c\u6ca1\u6709\u7f3a\u7684',
           '\u96f6\u4ef6\u6309\u7167\u5206\u7c7b\u6446\u597d\u4e86',
           '\u80fd\u8bf4\u51fa\u6bcf\u4e2a\u96f6\u4ef6\u662f\u5c0f\u8f66\u7684\u4ec0\u4e48',
           '\u62cd\u4e86\u5168\u5bb6\u798f\u7167\u7247'])
star('\u2b50 \u606d\u559c\uff01\u7b2c 0 \u5173\u901a\u8fc7\uff01\u8d34\u4e00\u9897\u661f\uff01')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  LEVEL 1 - LED
# ══════════════════════════════════════════════════

h('\U0001F4A1 \u7b2c 1 \u5173\uff1a\u70b9\u4eae\u5c0f\u706f\uff01', level=1)
p('\U0001F3AF \u4efb\u52a1\uff1a\u8ba9 LED \u706f\u4eae\u8d77\u6765\uff0c\u8ba9\u8702\u9e23\u5668\u54cd\u8d77\u6765\uff0c\u7528\u6309\u94ae\u63a7\u5236\u5b83\u4eec\uff01', bold=True, size=12)
p('\u8fd9\u662f\u4f60\u7b2c\u4e00\u6b21\u8ba9\u7535\u5b50\u96f6\u4ef6\u542c\u4f60\u7684\u8bdd\u2014\u2014\u8d85\u7ea7\u9177\uff01', italic=True)

h('\U0001F4E6 \u4ece\u96f6\u4ef6\u5806\u91cc\u62ff\u51fa\u6765', level=2)
table(
    ['\u96f6\u4ef6', '\u6570\u91cf'],
    [
        ['\U0001F9E0 ESP32 \u5f00\u53d1\u677f', '\xd71'],
        ['\U0001F49A \u7eff\u8272 LED', '\xd71'],
        ['\u2764\uFE0F \u7ea2\u8272 LED', '\xd71'],
        ['\U0001F514 \u8702\u9e23\u5668', '\xd71'],
        ['\U0001F518 \u6309\u94ae\u5f00\u5173', '\xd71'],
        ['\U0001F50C \u675c\u90a6\u7ebf', '\u51e0\u6839'],
        ['\U0001F50C USB-C \u6570\u636e\u7ebf', '\xd71\uff08\u8fde\u7535\u8111\u7528\uff09'],
    ]
)

h('\u7b2c \u2460 \u6b65\uff1a\u8ba4\u8bc6 ESP32 \u7684\u5c0f\u811a', level=2)
p('ESP32 \u4e24\u8fb9\u6709\u5f88\u591a\u5c0f\u811a\uff08\u5f15\u811a\uff09\uff0c\u6bcf\u53ea\u811a\u90fd\u6709\u4e00\u4e2a\u7f16\u53f7\u3002\u6211\u4eec\u8981\u7528\u8fd9\u4e9b\u811a\uff1a')
table(
    ['\u5f15\u811a\u7f16\u53f7', '\u63a5\u4ec0\u4e48'],
    [
        ['GPIO 2', '\U0001F49A \u7eff\u8272 LED'],
        ['GPIO 4', '\u2764\uFE0F \u7ea2\u8272 LED'],
        ['GPIO 5', '\U0001F518 \u6309\u94ae'],
        ['GPIO 23', '\U0001F514 \u8702\u9e23\u5668'],
        ['GND', '\u6240\u6709\u96f6\u4ef6\u7684\u8d1f\u6781\uff08-\uff09'],
    ]
)
p('\U0001F4A1 GND \u5c31\u662f\u5730\uff0c\u6240\u6709\u96f6\u4ef6\u90fd\u8981\u63a5\u4e00\u6839\u7ebf\u5230 GND', italic=True, size=10)

h('\u7b2c \u2461 \u6b65\uff1a\u63a5\u7ebf\uff01\u50cf\u62fc\u79ef\u6728\u4e00\u6837', level=2)
p('\u7528\u675c\u90a6\u7ebf\u628a\u96f6\u4ef6\u548c ESP32 \u8fde\u8d77\u6765\uff0c\u4e00\u6839\u7ebf\u4e00\u6839\u7ebf\u6765\uff0c\u522b\u6025\uff1a')

p('\U0001F7E2 \u63a5\u7eff\u8272 LED\uff1a', bold=True)
code('    ESP32 GPIO 2 \u2500\u2500 \u675c\u90a6\u7ebf \u2500\u2500 LED \u957f\u811a\uff08+\uff09\n'
     '                                LED \u77ed\u811a\uff08-\uff09\u2500\u2500 \u675c\u90a6\u7ebf \u2500\u2500 GND\n'
     '\n'
     '    \U0001F4A1 LED \u6709\u4e00\u957f\u4e00\u77ed\u4e24\u53ea\u811a\n'
     '       \u957f\u811a = \u6b63\u6781\uff08+\uff09= \u63a5 GPIO\n'
     '       \u77ed\u811a = \u8d1f\u6781\uff08-\uff09= \u63a5 GND')

p('\U0001F534 \u63a5\u7ea2\u8272 LED\uff1a\uff08\u540c\u4e0a\uff0c\u63a5 GPIO 4\uff09', bold=True)
p('\U0001F514 \u63a5\u8702\u9e23\u5668\uff1a', bold=True)
code('    ESP32 GPIO 23 \u2500\u2500 \u675c\u90a6\u7ebf \u2500\u2500 \u8702\u9e23\u5668 +\n'
     '    ESP32 GND     \u2500\u2500 \u675c\u90a6\u7ebf \u2500\u2500 \u8702\u9e23\u5668 -\n'
     '    \U0001F4A1 \u8702\u9e23\u5668\u4e0a\u9762\u5370\u7740 + \u548c - \uff0c\u522b\u63a5\u53cd\u4e86\uff01')

p('\U0001F518 \u63a5\u6309\u94ae\uff1a', bold=True)
code('    ESP32 GPIO 5 \u2500\u2500 \u675c\u90a6\u7ebf \u2500\u2500 \u6309\u94ae\u7684\u4e00\u53ea\u811a\n'
     '    ESP32 GND    \u2500\u2500 \u675c\u90a6\u7ebf \u2500\u2500 \u6309\u94ae\u7684\u53e6\u4e00\u811a')

h('\u7b2c \u2462 \u6b65\uff1a\u68c0\u67e5\u63a5\u7ebf', level=2)
checklist(['\u7eff\u8272 LED \u957f\u811a \u2192 GPIO 2',
           '\u7eff\u8272 LED \u77ed\u811a \u2192 GND',
           '\u7ea2\u8272 LED \u957f\u811a \u2192 GPIO 4',
           '\u7ea2\u8272 LED \u77ed\u811a \u2192 GND',
           '\u8702\u9e23\u5668\uff08+\uff09\u2192 GPIO 23',
           '\u8702\u9e23\u5668\uff08-\uff09\u2192 GND',
           '\u6309\u94ae\u4e00\u811a \u2192 GPIO 5',
           '\u6309\u94ae\u53e6\u4e00\u811a \u2192 GND'])

h('\u7b2c \u2463 \u6b65\uff1a\u63d2\u4e0a USB \u7ebf\uff0c\u8fde\u7535\u8111', level=2)
p('\u7528 USB-C \u7ebf\u628a ESP32 \u548c\u7535\u8111\u8fde\u8d77\u6765\u3002ESP32 \u4e0a\u4f1a\u4eae\u4e00\u4e2a\u5c0f\u7ea2\u706f = \u6709\u7535\u4e86\uff01')

h('\u7b2c \u2464 \u6b65\uff1a\u70e7\u5f55\u7a0b\u5e8f\uff01', level=2)
p('\u8ba9\u5927\u4eba\u5e2e\u4f60\u6253\u5f00 Arduino IDE\uff0c\u6253\u5f00\u6587\u4ef6\uff1a')
p('tests/test_1_led_button/test_1_led_button.ino', bold=True, size=10)
p('\u9009\u62e9\u5f00\u53d1\u677f ESP32 Dev Module\uff0c\u70b9\u4e0a\u4f20\u6309\u94ae \u25b6\uFE0F')

h('\u7b2c \u2465 \u6b65\uff1a\u73a9\u8d77\u6765\uff01', level=2)
p('\u4e0a\u4f20\u6210\u529f\u540e\uff1a')
bullet('\u5f00\u673a\u706f\u5149\u79c0\uff01\U0001F49A\u2764\uFE0F\U0001F49A\u2764\uFE0F\U0001F49A\u2764\uFE0F \u95ea\u4e09\u4e0b')
bullet('\u7136\u540e\u7ea2\u706f\u4eae\u7740 = \u5f85\u673a')
bullet('\u6309\u4e00\u4e0b\u6309\u94ae \U0001F518 \u2192 \U0001F49A \u7eff\u706f\u4eae\uff01\u201c\u5600\u201d\uff08\u9ad8\u97f3\uff09')
bullet('\u518d\u6309\u4e00\u4e0b \U0001F518 \u2192 \u2764\uFE0F \u7ea2\u706f\u4eae\uff01\u201c\u561f\u201d\uff08\u4f4e\u97f3\uff09')
bullet('\u4ea4\u66ff\u5207\u6362\uff01')

h('\U0001F3AE \u6311\u6218\u4efb\u52a1\uff08\u9009\u505a\uff0c\u989d\u5916\u52a0 \u2b50\uff09', level=3)
p('\u6253\u5f00\u4ee3\u7801\uff0c\u8bd5\u8bd5\u6539\u8fd9\u4e9b\u6570\u5b57\uff0c\u770b\u770b\u4f1a\u600e\u6837\uff1a')
table(
    ['\u627e\u5230\u8fd9\u884c\u4ee3\u7801', '\u6539\u6539\u770b', '\u4f1a\u600e\u6837\uff1f'],
    [
        ['tone(BUZZER_PIN, 1500, 200)', '\u628a 1500 \u6539\u6210 3000', '\u58f0\u97f3\u53d8___\u4e86\uff01'],
        ['tone(BUZZER_PIN, 800, 200)', '\u628a 800 \u6539\u6210 200', '\u58f0\u97f3\u53d8___\u4e86\uff01'],
        ['delay(300)', '\u628a 300 \u6539\u6210 100', '\u706f\u95ea\u5f97___\u4e86\uff01'],
    ]
)

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u5f00\u673a\u706f\u5149\u79c0\u6b63\u5e38\u95ea',
           '\u6309\u6309\u94ae\u80fd\u5207\u6362\u7ea2\u706f/\u7eff\u706f',
           '\u8702\u9e23\u5668\u4f1a\u54cd'])
star('\u2b50 \u7b2c 1 \u5173\u901a\u8fc7\uff01\u4f60\u5df2\u7ecf\u662f\u5c0f\u5c0f\u7a0b\u5e8f\u5458\u4e86\uff01')

doc.add_page_break()

# ══════════════════════════════════════════════════
#  LEVELS 2-10: Shorter format to keep the script manageable
# ══════════════════════════════════════════════════

# -- Level 2 --
h('\U0001F9BF \u7b2c 2 \u5173\uff1a\u8ba9\u5c0f\u8f66\u8d70\u8def\uff01', level=1)
p('\U0001F3AF \u4efb\u52a1\uff1a\u8ba9\u4e24\u4e2a\u8f6e\u5b50\u8f6c\u8d77\u6765\uff0c\u5c0f\u8f66\u80fd\u524d\u8fdb\u3001\u540e\u9000\u3001\u8f6c\u5f2f\uff01', bold=True, size=12)

h('\U0001F4E6 \u4f60\u9700\u8981', level=2)
table(['\u96f6\u4ef6', '\u6570\u91cf'],
      [['\U0001F9BF N20 \u51cf\u901f\u7535\u673a', '\xd72'],
       ['\U0001F534 \u6a61\u80f6\u8f6e\u5b50', '\xd72'],
       ['\U0001F4CB TB6612 \u7535\u673a\u9a71\u52a8\u677f', '\xd71']])

h('\u7b2c \u2460 \u6b65\uff1a\U0001F466 \u4f60\u6765\uff01\u628a\u8f6e\u5b50\u88c5\u5230\u7535\u673a\u4e0a\uff01', level=2)
p('\u628a\u6a61\u80f6\u8f6e\u5b50\u7528\u529b\u6309\u5230\u7535\u673a\u7684\u8f6c\u8f74\u4e0a\uff0c\u505a\u4e24\u7ec4\uff01')

h('\u7b2c \u2461 \u6b65\uff1a\U0001F468 \u5927\u4eba\u63a5\u7ebf', level=2)
table(['\u5f15\u811a', '\u63a5\u5230', '\u7528\u9014'],
      [['GPIO 25', 'TB6612 AIN1', '\u5de6\u7535\u673a\u65b9\u5411'],
       ['GPIO 26', 'TB6612 AIN2', '\u5de6\u7535\u673a\u65b9\u5411'],
       ['GPIO 32', 'TB6612 PWMA', '\u5de6\u7535\u673a\u901f\u5ea6'],
       ['GPIO 27', 'TB6612 BIN1', '\u53f3\u7535\u673a\u65b9\u5411'],
       ['GPIO 14', 'TB6612 BIN2', '\u53f3\u7535\u673a\u65b9\u5411'],
       ['GPIO 33', 'TB6612 PWMB', '\u53f3\u7535\u673a\u901f\u5ea6'],
       ['GPIO 13', 'TB6612 STBY', '\u9a71\u52a8\u677f\u5f00\u5173']])

h('\u7b2c \u2462 \u6b65\uff1a\u26a0\uFE0F \u5148\u67b6\u7a7a\uff01', level=2)
warn('\u26a0\uFE0F \u5b89\u5168\u63d0\u793a\uff1a\u628a\u5c0f\u8f66\u7ffb\u8fc7\u6765\uff01\u8f6e\u5b50\u671d\u4e0a\uff01\n\u5148\u8ba9\u8f6e\u5b50\u60ac\u7a7a\u8f6c\uff0c\u786e\u8ba4\u65b9\u5411\u5bf9\u4e86\u518d\u653e\u5230\u684c\u4e0a\uff01')

h('\u7b2c \u2463 \u6b65\uff1a\u70e7\u5f55\u6d4b\u8bd5', level=2)
p('tests/test_2_motors/test_2_motors.ino', bold=True, size=10)

h('\u7b2c \u2464 \u6b65\uff1a\U0001F466 \u6309\u6309\u94ae\uff0c\u6307\u6325\u5c0f\u8f66\uff01', level=2)
table(['\u6309\u7b2c\u51e0\u4e0b', '\u52a8\u4f5c', '\u8f6e\u5b50\u600e\u4e48\u8f6c'],
      [['\u7b2c 1 \u4e0b', '\U0001F53C \u524d\u8fdb', '\u4e24\u4e2a\u8f6e\u5b50\u90fd\u5f80\u524d\u8f6c'],
       ['\u7b2c 2 \u4e0b', '\U0001F53D \u540e\u9000', '\u4e24\u4e2a\u8f6e\u5b50\u90fd\u5f80\u540e\u8f6c'],
       ['\u7b2c 3 \u4e0b', '\u25c0\uFE0F \u5de6\u8f6c', '\u5de6\u8f6e\u540e\u9000+\u53f3\u8f6e\u524d\u8fdb'],
       ['\u7b2c 4 \u4e0b', '\u25b6\uFE0F \u53f3\u8f6c', '\u5de6\u8f6e\u524d\u8fdb+\u53f3\u8f6e\u540e\u9000'],
       ['\u7b2c 5 \u4e0b', '\u23f9\uFE0F \u505c', '\u90fd\u4e0d\u8f6c']])

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u8f6e\u5b50\u88c5\u597d\u4e86\uff08\u81ea\u5df1\u6309\u4e0a\u53bb\u7684\uff01\uff09',
           '\u524d\u8fdb\u3001\u540e\u9000\u65b9\u5411\u6b63\u786e',
           '\u5de6\u8f6c\u3001\u53f3\u8f6c\u65b9\u5411\u6b63\u786e'])
star('\u2b50 \u7b2c 2 \u5173\u901a\u8fc7\uff01\u5c0f\u8f66\u8fc8\u51fa\u4e86\u7b2c\u4e00\u6b65\uff01')
doc.add_page_break()

# -- Level 3 --
h('\U0001F441 \u7b2c 3 \u5173\uff1a\u88c5\u773c\u775b \u2014 \u518d\u4e5f\u4e0d\u6389\u684c\u5b50\uff01', level=1)
p('\U0001F3AF \u88c5 3 \u4e2a\u60ac\u5d16\u773c\u775b\uff0c\u5c0f\u8f66\u770b\u5230\u684c\u5b50\u8fb9\u7f18\u5c31\u540e\u9000\uff01', bold=True, size=12)
p('\u8fd9\u662f\u6700\u91cd\u8981\u7684\u5b89\u5168\u529f\u80fd\uff01', italic=True, color=RED_WARN)

h('\u539f\u7406', level=2)
code('  \u3010\u5b89\u5168\uff01\u6709\u684c\u9762\u3011              \u3010\u5371\u9669\uff01\u8981\u6389\u4e86\uff01\u3011\n'
     '\n'
     '     \u4f20\u611f\u5668                       \u4f20\u611f\u5668\n'
     '     \u2193\u2193\u2193 \u53d1\u51fa\u7ea2\u5916\u7ebf               \u2193\u2193\u2193 \u53d1\u51fa\u7ea2\u5916\u7ebf\n'
     '     \u2191\u2191\u2191 \u53cd\u5c04\u56de\u6765\u4e86\uff01             ........  \u5c04\u51fa\u53bb\uff0c\u6ca1\u56de\u6765\uff01\n'
     '   \u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501 \u684c\u9762                       \u2503 \u684c\u5b50\u8fb9\n'
     '\n'
     '     \u2192 \u7ee7\u7eed\u8d70\uff01                    \u2192 \u540e\u9000\uff01\u8f6c\u5f2f\uff01')

h('\u63a5\u7ebf', level=2)
table(['\u4f20\u611f\u5668', '\u63a5\u5230 ESP32'],
      [['\u5de6\u4f20\u611f\u5668 DO', 'GPIO 34'],
       ['\u4e2d\u4f20\u611f\u5668 DO', 'GPIO 35'],
       ['\u53f3\u4f20\u611f\u5668 DO', 'GPIO 36'],
       ['\u6240\u6709 VCC', '5V'],
       ['\u6240\u6709 GND', 'GND']])

h('\U0001F466 \u505a\u5b9e\u9a8c\uff01', level=2)
p('tests/test_3_cliff/test_3_cliff.ino', bold=True, size=10)
bullet('\u7528\u624b\u638c\u906e\u4f4f\u4f20\u611f\u5668 \u2192 \u5047\u88c5\u662f\u684c\u9762 \u2192 \U0001F49A \u7eff\u706f\u4eae')
bullet('\u628a\u624b\u62ff\u5f00 \u2192 \u5047\u88c5\u662f\u60ac\u5d16 \u2192 \u2764\uFE0F \u7ea2\u706f\u4eae')

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['3 \u4e2a\u4f20\u611f\u5668\u90fd\u80fd\u68c0\u6d4b',
           '\u5c0f\u8f66\u5230\u684c\u8fb9\u4f1a\u81ea\u52a8\u540e\u9000\u8f6c\u5f2f',
           '\u5728\u684c\u4e0a\u8dd1 1 \u5206\u949f\u4e0d\u6389\u4e0b\u53bb\uff01'])
star('\u2b50\u2b50 \u7b2c 3 \u5173\u901a\u8fc7\uff01\u53cc\u661f\uff01\u6700\u91cd\u8981\u7684\u5b89\u5168\u91cc\u7a0b\u7891\uff01')
p('\U0001F4F8 \u62cd\u4e2a\u89c6\u9891\uff01\u53d1\u7ed9\u540c\u5b66\u770b\u8d85\u9177\uff01', italic=True, size=10, align='center')
doc.add_page_break()

# -- Level 4 --
h('\U0001F4E1 \u7b2c 4 \u5173\uff1a\u88c5\u96f7\u8fbe \u2014 \u4e0d\u649e\u4e1c\u897f\uff01', level=1)
p('\U0001F3AF \u88c5\u8d85\u58f0\u6ce2\u4f20\u611f\u5668\uff0c\u8ba9\u5c0f\u8f66\u770b\u5230\u684c\u4e0a\u7684\u676f\u5b50\u548c\u4e66\u672c\uff01', bold=True, size=12)

table(['\u4f20\u611f\u5668\u5f15\u811a', '\u63a5\u5230 ESP32', '\u4f5c\u7528'],
      [['Trig', 'GPIO 18', '\u53d1\u8d85\u58f0\u6ce2'],
       ['Echo', 'GPIO 19', '\u542c\u56de\u58f0'],
       ['VCC', '5V', '\u4f9b\u7535'],
       ['GND', 'GND', '\u63a5\u5730']])

h('\U0001F466 \u73a9\u6d4b\u8ddd\u6e38\u620f\uff01', level=2)
p('tests/test_4_ultrasonic/test_4_ultrasonic.ino', bold=True, size=10)
bullet('\u5148\u7528\u773c\u775b\u731c\uff1a\u8fd9\u4e2a\u676f\u5b50\u79bb\u4f20\u611f\u5668\u5927\u6982\u591a\u8fdc\uff1f')
bullet('\u770b\u4e32\u53e3\u6253\u5370\u7684\u771f\u5b9e\u8ddd\u79bb \u2192 \u731c\u5bf9\u4e86\u5417\uff1f')

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u80fd\u7528\u8d85\u58f0\u6ce2\u6d4b\u51fa\u8ddd\u79bb', '\u5c0f\u8f66\u9047\u5230\u676f\u5b50\u4f1a\u7ed5\u5f00'])
star('\u2b50 \u7b2c 4 \u5173\u901a\u8fc7\uff01')
doc.add_page_break()

# -- Level 5 --
h('\U0001F50A \u7b2c 5 \u5173\uff1a\u88c5\u5634\u5df4 \u2014 \u5c0f\u8f66\u4f1a\u8bf4\u8bdd\uff01', level=1)
p('\U0001F3AF \u5f55\u5236\u53ef\u7231\u7684\u53f0\u8bcd\uff0c\u8ba9\u5c0f\u8f66\u7528\u4f60\u7684\u58f0\u97f3\u8bf4\u8bdd\uff01', bold=True, size=12)
p('\u2b50 \u8fd9\u4e00\u5173\u4f60\u662f\u7edd\u5bf9\u4e3b\u89d2\uff01\u53ea\u6709\u4f60\u80fd\u5f55\u81ea\u5df1\u7684\u58f0\u97f3\uff01', bold=True, color=GOLD)

h('\U0001F466 \u5f55\u53f0\u8bcd\uff01\u4f60\u7684\u72ec\u5bb6\u4efb\u52a1\uff01', level=2)
p('\u6253\u5f00\u624b\u673a\u5f55\u97f3 APP\uff0c\u7528\u6700\u53ef\u7231\u7684\u58f0\u97f3\u5f55\u4e0b\u9762\u7684\u8bdd\u3002\u4e00\u5171 20 \u6761\uff01')

table(['\u7f16\u53f7', '\u6587\u4ef6\u540d', '\u53f0\u8bcd'],
      [['1', '0001.mp3', '\u4f60\u597d\u5440\uff01\u6211\u662f\u684c\u9762\u5c0f\u7ba1\u5bb6\uff0c\u6309\u4e00\u4e0b\u6211\u5c31\u5e2e\u4f60\u6253\u626b\u54e6~'],
       ['2', '0002.mp3', '\u51fa\u53d1\u54af\uff01\u8ba9\u6211\u628a\u684c\u5b50\u64e6\u5f97\u4eae\u4eae\u7684\uff01'],
       ['3', '0003.mp3', '\u4ea4\u7ed9\u6211\u5427\uff01\u6a61\u76ae\u5c51\u5c51\uff0c\u6211\u6765\u6536\u62fe\u4f60\u4eec\u5566~'],
       ['4', '0004.mp3', '\u5f00\u5de5\u5f00\u5de5\uff01\u561f\u561f\u561f~'],
       ['5', '0005.mp3', '\u547c~\u597d\u591a\u788e\u5c51\u5440\uff0c\u4e0d\u8fc7\u96be\u4e0d\u5012\u6211\uff01'],
       ['6', '0006.mp3', '\u55e1\u55e1\u55e1~\u6211\u662f\u52e4\u52b3\u7684\u5c0f\u871c\u8702~'],
       ['7', '0007.mp3', '\u8fd9\u91cc\u597d\u810f\u54e6\uff0c\u8ba9\u6211\u591a\u626b\u626b~'],
       ['8', '0008.mp3', '\u54fc\u54fc\u54fc~\u6253\u626b\u536b\u751f\u771f\u5f00\u5fc3~'],
       ['9', '0009.mp3', '\u52a0\u6cb9\u52a0\u6cb9\uff0c\u9a6c\u4e0a\u5c31\u5e72\u51c0\u5566\uff01'],
       ['10', '0010.mp3', '\u54ce\u5440\uff01\u5dee\u70b9\u6389\u4e0b\u53bb\uff0c\u597d\u9669\u597d\u9669~'],
       ['11', '0011.mp3', '\u524d\u9762\u6ca1\u8def\u4e86\uff0c\u6211\u8f6c\u4e2a\u5f2f~'],
       ['12', '0012.mp3', '\u545c\u545c\uff0c\u684c\u5b50\u8fb9\u8fb9\u597d\u5413\u4eba\uff0c\u6211\u6362\u4e2a\u65b9\u5411\uff01'],
       ['13', '0013.mp3', '\u524d\u9762\u6709\u4e1c\u897f\u6321\u4f4f\u4e86\uff0c\u6211\u7ed5\u8fc7\u53bb~'],
       ['14', '0014.mp3', '\u54ce\u5466\uff0c\u649e\u5230\u4e1c\u897f\u4e86\uff0c\u6362\u6761\u8def\u8d70~'],
       ['15', '0015.mp3', '\u8fd9\u91cc\u6709\u4e2a\u5927\u5bb6\u4f19\uff0c\u6211\u8fd8\u662f\u8eb2\u5f00\u5427~'],
       ['16', '0016.mp3', '\u6253\u626b\u5b8c\u5566\uff01\u684c\u5b50\u5e72\u5e72\u51c0\u51c0\uff0c\u8bb0\u5f97\u5938\u5938\u6211\u54e6~'],
       ['17', '0017.mp3', '\u4efb\u52a1\u5b8c\u6210\uff01\u4f60\u7684\u684c\u9762\u5df2\u7ecf\u95ea\u95ea\u53d1\u4eae\u5566\uff01'],
       ['18', '0018.mp3', '\u547c~\u7d2f\u6b7b\u6211\u4e86\uff0c\u4e0d\u8fc7\u770b\u5230\u5e72\u51c0\u7684\u684c\u5b50\u597d\u5f00\u5fc3\uff01'],
       ['19', '0019.mp3', '\u597d\u7684\u597d\u7684\uff0c\u6211\u4f11\u606f\u4e00\u4e0b~'],
       ['20', '0020.mp3', '\u6211\u7684\u5c0f\u809a\u5b50\u597d\u50cf\u88c5\u6ee1\u4e86\uff0c\u5e2e\u6211\u5012\u4e00\u4e0b\u5783\u573e\u5427~']])

h('\u5f55\u97f3\u5c0f\u6280\u5de7', level=3)
bullet('\u58f0\u97f3\u5927\u4e00\u70b9\uff0c\u79bb\u624b\u673a\u8fd1\u4e00\u70b9')
bullet('\u7528\u6d3b\u6cfc\u53ef\u7231\u7684\u8bed\u6c14')
bullet('\u5f55\u5b8c\u542c\u4e00\u904d\uff0c\u4e0d\u6ee1\u610f\u5c31\u91cd\u5f55')
bullet('\u5728\u5b89\u9759\u7684\u73af\u5883\u5f55')

h('\U0001F466 \u628a\u5f55\u97f3\u653e\u8fdb SD \u5361', level=2)
code('  SD \u5361/\n  \u2514\u2500\u2500 mp3/\n      \u251c\u2500\u2500 0001.mp3\n      \u251c\u2500\u2500 0002.mp3\n      \u251c\u2500\u2500 \u2026\u2026\n      \u2514\u2500\u2500 0020.mp3')
warn('\u26a0\uFE0F \u6587\u4ef6\u540d\u5fc5\u987b\u662f 4 \u4f4d\u6570\u5b57\uff010001 \u4e0d\u662f 1\uff01\n\u26a0\uFE0F SD \u5361\u8981\u683c\u5f0f\u5316\u4e3a FAT32 \u683c\u5f0f\uff01')

h('\U0001F468 \u5927\u4eba\u63a5\u7ebf', level=2)
table(['DFPlayer \u5f15\u811a', '\u63a5\u5230', '\u8bf4\u660e'],
      [['VCC', '5V', '\u4f9b\u7535'],
       ['GND', 'GND', '\u63a5\u5730'],
       ['RX', 'GPIO 16\uff08\u7ecf 1k\u03a9 \u7535\u963b\uff09', '\u63a5\u6536\u6307\u4ee4'],
       ['TX', 'GPIO 17', '\u53d1\u9001\u72b6\u6001'],
       ['SPK_1', '\u55ad\u53ed +', '\u58f0\u97f3\u8f93\u51fa'],
       ['SPK_2', '\u55ad\u53ed -', '\u58f0\u97f3\u8f93\u51fa']])

h('\U0001F466 \u8bd5\u542c\uff01', level=2)
p('tests/test_5_dfplayer/test_5_dfplayer.ino', bold=True, size=10)
p('\u5b83\u4f1a\u81ea\u52a8\u4e00\u6761\u4e00\u6761\u64ad\u653e\u4f60\u5f55\u7684\u53f0\u8bcd\uff01\u6309\u6309\u94ae\u53ef\u4ee5\u8df3\u5230\u4e0b\u4e00\u6761\u3002')

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['20 \u6761\u53f0\u8bcd\u5168\u90e8\u5f55\u5b8c', 'SD \u5361\u91cc\u6587\u4ef6\u547d\u540d\u6b63\u786e',
           '\u55ad\u53ed\u80fd\u64ad\u653e\u51fa\u4f60\u7684\u58f0\u97f3', '\u97f3\u91cf\u5408\u9002'])
star('\u2b50\u2b50 \u7b2c 5 \u5173\u901a\u8fc7\uff01\u53cc\u661f\uff01\u4f60\u7ed9\u4e86\u5c0f\u8f66\u81ea\u5df1\u7684\u58f0\u97f3\uff01')
doc.add_page_break()

# -- Level 6 --
h('\U0001F3A4 \u7b2c 6 \u5173\uff1a\u88c5\u8033\u6735 \u2014 \u542c\u61c2\u4f60\u8bf4\u8bdd\uff01', level=1)
p('\U0001F3AF \u8ba9\u5c0f\u8f66\u80fd\u542c\u61c2\u300c\u6253\u626b\u300d\u300c\u505c\u4e0b\u300d\u300c\u4f60\u597d\u300d\uff01', bold=True, size=12)

table(['LU-ASR01 \u5f15\u811a', '\u63a5\u5230 ESP32'],
      [['VCC', '5V'], ['GND', 'GND'], ['TXD', 'GPIO 21'], ['RXD', 'GPIO 22']])

h('\U0001F466 \u8bbe\u7f6e\u5173\u952e\u8bcd\uff01\u62d6\u79ef\u6728\uff01', level=2)
p('\u5728 TIANWEN BLOCK \u91cc\u62d6\u62fd\u79ef\u6728\uff1a')
table(['\u542c\u5230\u4ec0\u4e48', '\u53d1\u9001\u4ec0\u4e48', '\u5c0f\u8f66\u505a\u4ec0\u4e48'],
      [['\u300c\u6253\u626b\u300d\u6216\u300c\u5f00\u59cb\u300d', '0x01', '\u5f00\u59cb\u6e05\u626b'],
       ['\u300c\u505c\u4e0b\u300d\u6216\u300c\u505c\u300d', '0x02', '\u505c\u6b62\u6e05\u626b'],
       ['\u300c\u4f60\u597d\u300d', '0x03', '\u56de\u590d\u6253\u62db\u547c\uff08\u5f69\u86cb\uff09']])

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u300c\u6253\u626b\u300d\u80fd\u88ab\u8bc6\u522b', '\u300c\u505c\u4e0b\u300d\u80fd\u88ab\u8bc6\u522b', '\u300c\u4f60\u597d\u300d\u80fd\u88ab\u8bc6\u522b'])
star('\u2b50 \u7b2c 6 \u5173\u901a\u8fc7\uff01\u4e0d\u7528\u6309\u6309\u94ae\uff0c\u8bf4\u8bdd\u5c31\u80fd\u63a7\u5236\uff01')
doc.add_page_break()

# -- Level 7 --
h('\U0001F9F9 \u7b2c 7 \u5173\uff1a\u88c5\u626b\u628a \u2014 \u771f\u7684\u80fd\u626b\u5730\uff01', level=1)
p('\U0001F3AF \u88c5\u4e0a\u65cb\u8f6c\u6bdb\u5237 + \u5438\u5c18\u98ce\u6247\uff0c\u8ba9\u5c0f\u8f66\u771f\u6b63\u80fd\u6e05\u626b\u684c\u9762\uff01', bold=True, size=12)

h('\U0001F466 DIY \u6bdb\u5237\uff01', level=2)
bullet('\u627e\u4e00\u4e2a\u74f6\u76d6\u6216\u5c0f\u5706\u7ba1')
bullet('\u7528\u70ed\u7194\u80f6\u7c98\u4e0a\u77ed\u5237\u6bdb')
bullet('\u88c5\u5230 130 \u7535\u673a\u7684\u8f6c\u8f74\u4e0a')

table(['\u90e8\u4ef6', '\u901a\u8fc7', '\u63a7\u5236\u5f15\u811a'],
      [['130 \u7535\u673a\uff08\u6bdb\u5237\uff09', 'MOS \u7ba1\u6a21\u5757 1', 'GPIO 12'],
       ['3010 \u98ce\u6247', 'MOS \u7ba1\u6a21\u5757 2', 'GPIO 15']])

h('\U0001F466 \u6e05\u626b\u6548\u679c\u5927\u6d4b\u8bd5\uff01', level=2)
p('\u5728\u684c\u4e0a\u6492\u4e00\u4e9b\u6a61\u76ae\u5c51\uff0c\u8ba9\u5c0f\u8f66\u5f00\u8fc7\u53bb\uff0c\u770b\u80fd\u4e0d\u80fd\u626b\u5e72\u51c0\uff01')

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u6bdb\u5237\u80fd\u65cb\u8f6c', '\u98ce\u6247\u80fd\u5439\u98ce', '\u80fd\u626b\u8d77\u6a61\u76ae\u5c51', '\u788e\u5c51\u8fdb\u4e86\u6536\u96c6\u76d2'])
star('\u2b50 \u7b2c 7 \u5173\u901a\u8fc7\uff01\u5c0f\u8f66\u771f\u7684\u4f1a\u626b\u5730\u4e86\uff01')
doc.add_page_break()

# -- Level 8 --
h('\U0001F50B \u7b2c 8 \u5173\uff1a\u88c5\u7535\u6c60 \u2014 \u81ea\u7531\u5954\u8dd1\uff01', level=1)
p('\U0001F3AF \u88c5\u4e0a\u7535\u6c60\uff0c\u62d4\u6389 USB \u7ebf\uff0c\u5c0f\u8f66\u771f\u6b63\u65e0\u7ebf\u4e86\uff01', bold=True, size=12)

warn('\u26a0\uFE0F \u7535\u6c60\u5b89\u5168\u89c4\u5219\uff08\u5f88\u91cd\u8981\uff01\uff09\n\n1. \u63a5\u7ebf\u65f6\u7535\u6c60\u76d2\u5f00\u5173\u8981\u5173\u7740\uff01\n2. \u770b\u6e05\u6b63\u6781(+)\u548c\u8d1f\u6781(-)\u518d\u63a5\uff01\n3. \u4e0d\u8981\u7528\u91d1\u5c5e\u7ebf\u628a\u7535\u6c60\u4e24\u5934\u76f4\u63a5\u8fde\u5728\u4e00\u8d77\uff01\n4. \u7535\u6c60\u4e0d\u8981\u6254\u8fdb\u706b\u91cc\uff01')

code('  \u7535\u6c601 + \u7535\u6c602 \u4e32\u8054 = 3.7V + 3.7V = 7.4V\n'
     '       \u2502\n'
     '    [\u5f00\u5173]\n'
     '       \u2502\n'
     '  [\u964d\u538b\u6a21\u5757] \u2192 7.4V \u53d8\u6210 5V\n'
     '       \u2502\n'
     '       \u251c\u2500\u2500\u2192 ESP32 \u5927\u8111\n'
     '       \u251c\u2500\u2500\u2192 DFPlayer \u5634\u5df4\n'
     '       \u2514\u2500\u2500\u2192 \u4f20\u611f\u5668\u4eec')

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u4e0d\u63d2 USB \u7ebf\uff0c\u5c0f\u8f66\u4e5f\u80fd\u5de5\u4f5c', '\u706f\u4eae\u3001\u55ad\u53ed\u54cd\u3001\u8f6e\u5b50\u8f6c'])
star('\u2b50 \u7b2c 8 \u5173\u901a\u8fc7\uff01\u5c0f\u8f66\u81ea\u7531\u4e86\uff01')
doc.add_page_break()

# -- Level 9 --
h('\U0001F3A8 \u7b2c 9 \u5173\uff1a\u7a7f\u8863\u670d \u2014 \u53d8\u6f02\u4eae\uff01', level=1)
p('\U0001F3AF \u7ed9\u5c0f\u8f66\u8bbe\u8ba1\u5916\u89c2\uff0c\u8ba9\u5b83\u53d8\u6210\u4f60\u60f3\u8981\u7684\u6837\u5b50\uff01', bold=True, size=12)
p('\u8fd9\u4e00\u5173\u5b8c\u5168\u7531\u4f60\u505a\u4e3b\uff01', bold=True, color=GOLD)

h('\U0001F466 \u4f60\u6765\u8bbe\u8ba1\uff01', level=2)
p('\u62ff\u4e00\u5f20\u7eb8\uff0c\u753b\u51fa\u4f60\u60f3\u8981\u7684\u5c0f\u8f66\u5916\u5f62\u2014\u2014\u5b83\u53ef\u4ee5\u662f\u4e00\u53ea\u732b\u3001\u4e00\u67b6\u98de\u8239\u3001\u4e00\u4e2a\u673a\u5668\u4eba\u2026\u2026\u968f\u4f60\uff01')
bullet('\u7528\u5f69\u7b14\u753b\u56fe\u6848')
bullet('\u8d34\u8d34\u7eb8')
bullet('\u7ed9\u5c0f\u8f66\u53d6\u540d\u5b57')
bullet('\u753b\u4e24\u53ea\u5927\u773c\u775b')

h('\u7ec4\u88c5\u56fa\u5b9a', level=2)
checklist(['\u7535\u673a + \u8f6e\u5b50 \u2192 \u70ed\u7194\u80f6/\u87ba\u4e1d\u56fa\u5b9a',
           '\u4f20\u611f\u5668 \u2192 \u70ed\u7194\u80f6\u7c98\u597d',
           'ESP32 + \u9a71\u52a8\u677f \u2192 \u87ba\u4e1d+\u94dc\u67f1\u56fa\u5b9a',
           '\u7ebf\u7f06 \u2192 \u624e\u5e26\u624e\u6574\u9f50',
           '\u76d6\u4e0a\u5916\u58f3\uff01'])

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u5916\u58f3\u8bbe\u8ba1\u597d\u4e86', '\u6240\u6709\u96f6\u4ef6\u56fa\u5b9a\u7262\u4e86', '\u5c0f\u8f66\u53d8\u597d\u770b\u4e86\uff01'])
star('\u2b50 \u7b2c 9 \u5173\u901a\u8fc7\uff01')
doc.add_page_break()

# -- Level 10 --
h('\U0001F3C6 \u7b2c 10 \u5173\uff1a\u6700\u7ec8\u5f69\u6392\uff01', level=1)
p('\U0001F3AF \u7ec3\u4e60\u5c55\u793a\u6d41\u7a0b\uff0c\u51c6\u5907\u597d\u4e0a\u53f0\uff01', bold=True, size=12)

h('\U0001F3AC \u8868\u6f14\u6d41\u7a0b\uff08\u7ec3 3 \u904d\uff01\uff09', level=2)
table(['\u6b65\u9aa4', '\u4f60\u505a\u4ec0\u4e48', '\u5c0f\u8f66\u505a\u4ec0\u4e48'],
      [['1', '\u6253\u5f00\u7535\u6e90\u5f00\u5173', '\U0001F50A \u4f60\u597d\u5440\uff01\u6211\u662f\u684c\u9762\u5c0f\u7ba1\u5bb6~'],
       ['2', '\u5728\u684c\u4e0a\u6492\u6a61\u76ae\u5c51', '\uff08\u8ba9\u8bc4\u59d4\u770b\u5230\u684c\u9762\u662f\u810f\u7684\uff09'],
       ['3', '\u8bf4\u300c\u4f60\u597d\uff01\u300d', '\U0001F50A \u56de\u7b54\u6253\u62db\u547c\uff08\u5f69\u86cb\uff01\uff09'],
       ['4', '\u8bf4\u300c\u6253\u626b\uff01\u300d', '\U0001F49A\u7eff\u706f\u4eae + \u5f00\u59cb\u8dd1 + \u8bf4\u8bdd + \u626b\u788e\u5c51'],
       ['5', '\u770b\u5c0f\u8f66\u8868\u6f14', '\u5230\u684c\u8fb9\u540e\u9000 + \u7ed5\u5f00\u676f\u5b50 + \u8fb9\u626b\u8fb9\u8bf4\u8bdd'],
       ['6', '\u8bf4\u300c\u505c\u4e0b\uff01\u300d', '\u2764\uFE0F\u7ea2\u706f\u4eae + \u505c\u4e0b + \u6211\u4f11\u606f\u4e00\u4e0b~'],
       ['7', '\u62c9\u51fa\u6536\u96c6\u76d2', '\u5c55\u793a\u91cc\u9762\u7684\u788e\u5c51\uff01']])

h('\U0001F3A4 \u8bc4\u59d4\u53ef\u80fd\u95ee\u7684\u95ee\u9898', level=2)

p('Q: \u8fd9\u662f\u4ec0\u4e48\uff1f', bold=True)
p('A: \u8fd9\u662f\u6211\u505a\u7684\u667a\u80fd\u684c\u9762\u6e05\u6d01\u5c0f\u8f66\uff01\u505a\u5b8c\u4f5c\u4e1a\u684c\u9762\u6709\u5f88\u591a\u6a61\u76ae\u5c51\uff0c\u6211\u60f3\u505a\u4e00\u4e2a\u80fd\u81ea\u52a8\u5e2e\u6211\u6253\u626b\u7684\u5c0f\u8f66\u3002')

p('Q: \u5b83\u600e\u4e48\u77e5\u9053\u684c\u5b50\u8fb9\u5728\u54ea\uff1f', bold=True)
p('A: \u5b83\u809a\u5b50\u5e95\u4e0b\u6709 3 \u4e2a\u4f20\u611f\u5668\uff0c\u7528\u770b\u4e0d\u89c1\u7684\u7ea2\u5916\u7ebf\u5f80\u4e0b\u7167\uff0c\u5982\u679c\u5149\u6ca1\u53cd\u5c04\u56de\u6765\uff0c\u5c31\u77e5\u9053\u5230\u684c\u5b50\u8fb9\u4e86\u3002')

p('Q: \u4f60\u7528\u4e86\u4ec0\u4e48\u6280\u672f\uff1f', bold=True)
p('A: \u8bed\u97f3\u8bc6\u522b\u3001\u8bed\u97f3\u64ad\u62a5\u3001\u8d85\u58f0\u6ce2\u907f\u969c\u3001\u7ea2\u5916\u9632\u5760\u3001\u7f16\u7a0b\u30013D \u6253\u5370\u3001\u624b\u5de5\u7ec4\u88c5\u3002')

p('Q: \u54ea\u4e9b\u662f\u4f60\u81ea\u5df1\u505a\u7684\uff1f', bold=True)
p('A: \u6211\u81ea\u5df1\u5f55\u4e86\u6240\u6709\u7684\u8bed\u97f3\u53f0\u8bcd\u3001\u8bbe\u8ba1\u4e86\u5916\u58f3\u3001\u53c2\u4e0e\u4e86\u63a5\u7ebf\u548c\u8c03\u8bd5\u3001\u505a\u4e86\u6bdb\u5237\u3002\u7f16\u7a0b\u662f\u548c\u5927\u4eba\u4e00\u8d77\u5b66\u7740\u505a\u7684\u3002')

doc.add_paragraph()
p('\u2705 \u901a\u5173\u6761\u4ef6', bold=True, size=12, color=GREEN)
checklist(['\u8868\u6f14\u6d41\u7a0b\u7ec3\u4e86 3 \u904d\u4ee5\u4e0a', '\u80fd\u56de\u7b54\u8bc4\u59d4\u7684\u95ee\u9898', '\u5c0f\u8f66\u7a33\u5b9a\u8fd0\u884c\u4e0d\u51fa\u9519'])
star('\u2b50 \u7b2c 10 \u5173\u901a\u8fc7\uff01')
doc.add_page_break()

# ══════════════════════════════════════════════════
#  ENDING
# ══════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p('\U0001F389\U0001F389\U0001F389', size=36, align='center')
doc.add_paragraph()
p('\U0001F3C6 \u5168\u90e8\u901a\u5173\uff01\u606d\u559c\u4f60\uff01\U0001F3C6', bold=True, size=26, color=GOLD, align='center')
doc.add_paragraph()
p('\u4f60\u4ece\u4e00\u5806\u96f6\u4ef6\u5f00\u59cb', size=14, align='center')
p('\u4e00\u6b65\u4e00\u6b65\u628a\u5b83\u4eec\u53d8\u6210\u4e86', size=14, align='center')
p('\u4e00\u8f86\u4f1a\u8bf4\u8bdd\u3001\u4f1a\u626b\u5730\u3001\u4f1a\u907f\u969c\u7684\u667a\u80fd\u5c0f\u8f66\uff01', bold=True, size=16, color=BLUE, align='center')

doc.add_paragraph()
p('\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50\u2b50', size=16, align='center')
p('\u6536\u96c6\u4e86 12 \u9897\u661f\uff01', bold=True, size=14, align='center')

doc.add_paragraph()
p('\u4f60\u5b66\u4f1a\u4e86\uff1a', bold=True, size=13)
for item in ['\u8ba4\u8bc6\u7535\u5b50\u96f6\u4ef6', 'LED\u3001\u8702\u9e23\u5668\u3001\u6309\u94ae\u63a7\u5236',
             '\u7535\u673a\u63a7\u5236\u548c\u5dee\u901f\u8f6c\u5411', '\u7ea2\u5916\u4f20\u611f\u5668\u539f\u7406',
             '\u8d85\u58f0\u6ce2\u6d4b\u8ddd', '\u5f55\u97f3\u548c SD \u5361\u4f7f\u7528',
             '\u8bed\u97f3\u8bc6\u522b\u914d\u7f6e', '\u673a\u68b0\u7ed3\u6784\u7ec4\u88c5',
             '\u4f9b\u7535\u7cfb\u7edf\u642d\u5efa', '\u4ea7\u54c1\u8bbe\u8ba1\u548c\u5c55\u793a']:
    bullet('\u2705 ' + item)

doc.add_paragraph()
doc.add_paragraph()
p('\u4f60\u5df2\u7ecf\u662f\u4e00\u4e2a\u771f\u6b63\u7684\u5c0f\u521b\u5ba2\u4e86\uff01\U0001F680', bold=True, size=18, color=BLUE, align='center')

# ── Save ──

output = '/Users/willas/src/maker-project/desk-cleaner/\u5c0f\u670b\u53cb\u7684\u7ec4\u88c5\u624b\u518c.docx'
doc.save(output)
print('Done: ' + output)
